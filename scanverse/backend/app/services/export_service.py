"""
Export service: turns a Document's processed page images (+ OCR text) into
downloadable PDF, DOCX, or TXT files.
"""

from __future__ import annotations

import os

import fitz  # PyMuPDF
from PIL import Image
from docx import Document as DocxDocument
from docx.shared import Inches

from app.core.config import settings


def build_pdf(image_paths: list[str], output_path: str) -> str:
    """Combine a list of image files into a single multi-page PDF."""
    doc = fitz.open()
    for path in image_paths:
        img = Image.open(path).convert("RGB")
        width_pt, height_pt = img.width * 72 / 200, img.height * 72 / 200  # assume ~200dpi scans
        page = doc.new_page(width=width_pt, height=height_pt)
        rect = fitz.Rect(0, 0, width_pt, height_pt)
        page.insert_image(rect, filename=path)
    doc.save(output_path)
    doc.close()
    return output_path


def build_docx(image_paths: list[str], ocr_texts: list[str | None], output_path: str) -> str:
    """Create a Word document with each page's image followed by its OCR text (if any)."""
    doc = DocxDocument()
    for i, path in enumerate(image_paths):
        doc.add_picture(path, width=Inches(6))
        text = ocr_texts[i] if i < len(ocr_texts) else None
        if text:
            doc.add_paragraph(text)
        if i < len(image_paths) - 1:
            doc.add_page_break()
    doc.save(output_path)
    return output_path


def build_txt(ocr_texts: list[str | None], output_path: str) -> str:
    with open(output_path, "w", encoding="utf-8") as f:
        content = "\n\n--- Page Break ---\n\n".join(t or "" for t in ocr_texts)
        f.write(content)
    return output_path


def ensure_export_dir() -> str:
    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    return settings.EXPORT_DIR
